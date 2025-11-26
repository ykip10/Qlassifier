""" Contains object definitions for PDF and word parsers, which parses 
documents into trees with certain assumptions on the structure of the documents. 
Supports parsing of: 

PDF documents which: 
    - Follow VCAA past examination formatting. May or may not be VCAA exams. 
Word documents which: 
    - Use header fonts for their headers. 
"""
import json
import re
from abc import ABC, abstractmethod
from pathlib import Path

import pymupdf
from docx import Document

from src.Qlassifier.trees import Tree
from src.Qlassifier import pdf_utils


class Parser(ABC):
    @abstractmethod
    def split_headings(self) -> Tree:
        """ Main method, splitting document into trees based off of headings. """
        pass


class PDFParser(Parser):
    """ Examination/VCAA report PDF parsing object. Supports cropping as well 
    as parsing PDF's based on Bold text, which are assumed to be headers, 
    which can then be accessed as a tree object. 
    """
    def __init__(self, path: str, cr_coords: tuple[int] = (0.07, 0.04, 0.93, 0.85)):
        """ Initialise PDF parsing object. 

        path     : Path of pdf to be parsed.
        cr_coords: Controls how much we crop in parsing. The cropbox we apply
                   will have PDF coordinates equal to the element-wise
                   multiplication of this parameter with (w, h, w, h), where
                   w and h are the width and height of the PDF respectively. 
        """
        self.path = path
        self.cr_coords = cr_coords 
        self._doc = self.load_doc()
        self._copies = []           # keep track of copies to close
        self.qn_hierarchies_regex = [       # question splits
                r"^S(?i:ection) [A-Z]", 	# level 1
                r"Question \d+",			# level 2
                r"[a-h]\.",					# level 3
                r"i+\.",					# level 4
        ]

    @property
    def doc(self):
        """ Returns a copy of the document. """
        doc_bytes = self._doc.write()
        new_doc = pymupdf.open(stream=doc_bytes,filetype="pdf")
        self._copies.append(new_doc)
        return new_doc

    def load_doc(self) -> pymupdf.Document:
        """ Loads the document. """
        try:
            doc = pymupdf.open(self.path)
        except Exception: 
            print(f"Error loading word document \'{self.path}\'")
            doc = None
        return doc

    def close(self):
        """ Closes PDF and all copies of it. """
        if self._doc is not None:
            self._doc.close()
        # Close all copies 
        for doc in self.copies:
            doc.close()
        self._copies.clear()

    def split_headings(self) -> Tree:
        """ Extracts text from a PDF using PyMuPdf and sorts them into Sections and Questions.

        Goes through bolded text as candidates for splitting, then checks if the bolded text are 
        valid question splits. Sorts extracted text into a Tree object and returns it. 

        The assumption is that new questions are labelled according to predefined QN_HIERARCHIES. 
        """
        # Find the last page, then crop the pdf. 
        final_page = self._find_final_page()
        doc = self.crop()

        curr = None
        nodes = []
        for page_idx in range(final_page+1): 
            page = doc[page_idx]
            txt_props_lst = self._get_text_and_style(
                json.loads(page.get_text("json"))
            )

            # Extract questions and their text from this page 
            for text, is_bold, y0 in txt_props_lst:
                if is_bold:
                    # possible question split candidate
                    matches = [bool(re.match(pattern, text)) 
                               for pattern in self.qn_hierarchies_regex]
                    if any(matches):
                        # Initialise new node
                        curr = Tree(label=text, level=matches.index(True)+1,
                                    page_idx=page_idx)
                        nodes.append((curr, page_idx, y0))
                elif curr:
                    # ordinary text
                    mark_match = re.search(r"(\d+)\s*marks?", text)
                    if mark_match and not curr.marks:
                        curr.marks = int(mark_match.group(1))
                    curr.text += " " + text


        # sort by page then vertical height to get correct ordering  
        nodes.sort(key=lambda t: (t[1], t[2]))  
        nodes = [text for text, page_num, y0 in nodes] 

        root = Tree(label="root", level=0)
        root.build(nodes)
        return root

    def crop(self, save_path: str = "") -> pymupdf.Document:
        return pdf_utils.crop(self.doc, self.cr_coords, save_path)

    def _span_is_bold(self, span: dict) -> bool:
        """ Checks if a PyMuPdf span is bold or not. """
        flags = span.get("flags", 0) or 0
        font = (span.get("font") or "").lower()

        # Check for LaTeX symbols, which might be misclassified as bold. 
        if any(k in font for k in ["cm", "cmsy", "cmex", "stix", "symbol"]):
            return False
        
        return (flags & 16) != 0 or ("bold" in font)

    def _get_text_and_style(self, page_json: dict) -> tuple[str, bool, int]:
        """ Given a parsed page.get_text("json") dict, return a list of 
        (text, is_bold, y0) tuples, where y0 is the texts height on the page. 
        """
        out = []
        for block in page_json.get("blocks", []):
            # text blocks have type == 0
            if block.get("type", 0) != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    bold = self._span_is_bold(span)
                    y0 = span.get("bbox", [0, 0, 0, 0])[1]  # top y-coordinate
                    out.append((text, bold, y0))
        return out
    
    def _find_final_page(self) -> int:
        """ Find the index of the final examination page by scanning through the footers 
        and looking for an 'end of examination' flag. Input should be uncropped. 
        """
        for page_idx in range(len(self.doc)):
            # First, crop into only the footer portion of the page
            page = self.doc[page_idx]
            w, h = page.rect.width, page.rect.height
            new_rect = pymupdf.Rect(0, self.cr_coords[3]*h, w, h)
            page.set_cropbox(new_rect)

            # Search for 'end of examination flag' in footer
            footer_props_lst = self._get_text_and_style(
                json.loads(page.get_text("json"))
            )
            
            for text, _, _, in footer_props_lst:
                pattern = r"(?i)^end of (?:\w+\s+)*"+\
                          r"(questions?|answer|book|examination)\b"
                if re.match(pattern, text):
                    return page_idx
                
        return page_idx


class WordParser(Parser):
    """ Word parser object. Supports the splitting of word 
    documents into tree objects based on headers. 
    """
    def __init__(self, path: str):
        """ Initialises word parsing object. 
        Takes a path argument and loads in the word doc.
        """
        self.path = path
        self.doc = self._load_doc()

    def split_headings(self) -> Tree: 
        """ Return a list of (level, text) for headings found in the document 
        found at path. It operates on the assumption that the document's 
        headings are stylised as headings (as opposed to something like 
        boldness/font size).
        """
        headings = []
        # First, build the headings
        curr = None
        for page_idx in range(len(self.doc.paragraphs)):
            page = self.doc.paragraphs[page_idx]
            style_name = getattr(page.style, "name", "") or ""
            text = page.text.strip()
            if not text:
                continue
            
            # Sometimes, VCAA documents have headings not stylised as such (annoying!)
            # so we have to check for this and adjust logic accordingly
            qn_re = r"Question \d+"
            sn_re = r"Section [A-Z]"
            qn_match = re.match(qn_re, text)
            section_match = re.match(sn_re, text)

            unstylised_heading = "heading" not in style_name.lower() and \
                         (qn_match or section_match)
            
            if unstylised_heading:
                # This problem has only been found in report docs,
                # where the following values make sense. Band-aid solution,
                # more than anything.

                # infer question/section levels from previously seen headings, if possible
                section_level = next(
                    (h.level for h in headings if re.match(sn_re, getattr(h, "label", ""))),
                    None
                )
                if section_level is not None:
                    qn_level = section_level + 1
                else:
                    qn_level = next(
                        (h.level for h in headings if re.match(qn_re, getattr(h, "label", ""))),
                        None
                    )

                    # Nothing to infer from previous headings.
                    # Fallback to standard values for a sectioned document. 
                    # If not sectioned, this will be incorrect.  
                    if qn_level is None:
                        qn_level = 3
                        section_level = 2
                    else:
                        section_level = qn_level - 1

                level = qn_level if qn_match else section_level
                curr = Tree(label=text, level=level, page_idx=page_idx)
                headings.append(curr)
                continue
                
            if "heading" in style_name.lower():
                # Found heading
                parts = style_name.split()
                for part in parts[::-1]:
                    # Extract number (level)
                    if part.isdigit():
                        level = int(part)
                        break
                curr = Tree(label=text, level=level,
                            page_idx=page_idx)
                headings.append(curr)
            elif curr: 
                curr.text += " " + text
        # Build the tree then return
        root = Tree(label="root", level=0)
        root.build(headings)
        return root

    def _load_doc(self) -> Document:
        """ Loads the word document. """
        try:
            doc = Document(self.path)
        except Exception as e: 
            print(f"Error loading word document \'{self.path}\'."
                  f"Path probably does not point to an existent file.")
            raise e
        return doc 


class AutoParser:
    """ Automatically determines whether to use a 
    WordParser or PDF parser based on input path filetype.
    Does not allow customisation of cropbox for PDF parsing."""
    def __init__(self, path: str):
        if Path(path).suffix not in [".pdf", ".docx"]:
            raise ValueError("Unsupported file type. Only parses word documents/PDFs.")
        
        self.parser = WordParser(path) if Path(path).suffix == ".docx" else PDFParser(path)
